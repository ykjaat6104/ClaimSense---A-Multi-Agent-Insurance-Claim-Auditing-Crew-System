from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.db.models import Claim, derive_flags_and_insights


def build_docx_report(claim: Claim, dest: Path) -> str:
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("ClaimSense — Claim Decision Support Report", level=0)
    for run in title.runs:
        run.font.size = Pt(18)

    meta = doc.add_table(rows=6, cols=2)
    meta.style = "Light Shading Accent 1"
    meta_data = [
        ("Claim ID", str(claim.id)),
        ("Status", claim.status or "N/A"),
        ("Decision", claim.decision or "N/A"),
        ("Risk score", str(claim.risk_score or "N/A")),
        ("Fraud probability", f"{claim.fraud_probability}%" if claim.fraud_probability is not None else "N/A"),
        ("Adjuster action", (claim.adjuster_action or "N/A").replace("ADJUSTER_", "")),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v

    doc.add_heading("Approve argument", level=2)
    doc.add_paragraph(claim.approve_argument or "No argument provided.")

    doc.add_heading("Reject / risk argument", level=2)
    doc.add_paragraph(claim.reject_argument or "No argument provided.")

    doc.add_heading("Mediator output", level=2)
    doc.add_paragraph(json.dumps(claim.mediator_output or {}, indent=2))

    if claim.rag_chunks:
        doc.add_heading("Retrieved policy clauses (RAG)", level=2)
        for i, chunk in enumerate(claim.rag_chunks[:12], 1):
            doc.add_paragraph(f"{i}. {chunk}")

    flags, insights = derive_flags_and_insights(claim)
    if flags:
        doc.add_heading("Flags", level=2)
        for flag in flags:
            doc.add_paragraph(flag, style="List Bullet")
    if insights:
        doc.add_heading("Insights", level=2)
        for insight in insights:
            doc.add_paragraph(insight)

    doc.save(str(dest))
    return str(dest.resolve())
